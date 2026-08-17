"""Export a rendered CV to PDF and Word — both derived from the same source.

PDF: headless Chromium prints the generated HTML (same engine that produced the
original cv.pdf), so the PDF matches the print stylesheet exactly.

Word: a clean, editable .docx built from the model. It does not pixel-match the
HTML (recruiters want editable text), but it is generated from data/*.csv, so it
stays in sync with everything else.
"""
from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from .labels import DEFAULT_LANG, labels as _labels

CHROMIUM_CANDIDATES = ["chromium", "chromium-browser", "google-chrome", "chrome"]


def _chromium() -> str:
    for c in CHROMIUM_CANDIDATES:
        if shutil.which(c):
            return c
    raise RuntimeError("no chromium/chrome binary found for PDF export")


def to_pdf(html_path: str | Path, pdf_path: str | Path) -> Path:
    html_path, pdf_path = Path(html_path).resolve(), Path(pdf_path).resolve()
    with tempfile.TemporaryDirectory() as profile:
        cmd = [
            _chromium(), "--headless", "--no-sandbox", "--disable-gpu",
            f"--user-data-dir={profile}", "--no-pdf-header-footer",
            "--virtual-time-budget=8000",
            f"--print-to-pdf={pdf_path}", html_path.as_uri(),
        ]
        # Chromium emits harmless dbus/AppArmor warnings to stderr and a non-zero
        # exit even on success, so success is judged by the output file, not rc.
        subprocess.run(cmd, capture_output=True, timeout=120)
    if not pdf_path.exists() or pdf_path.stat().st_size < 1000:
        raise RuntimeError(f"PDF export produced no usable file at {pdf_path}")
    return pdf_path


def to_docx(data: dict, path: str | Path, profile: dict | None = None,
            lang: str = DEFAULT_LANG) -> Path:
    """Editable Word export — same source, same language as the page it sits next to.

    Every heading comes from the label pack. Hard-coded German headings would
    have shipped a Word file titled "Projekthistorie" from the English CV: the
    same defect the English page itself had, one layer further down where nobody
    reads it before sending it to a recruiter.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    L = _labels(lang)
    profile = profile or {}
    person = data["person"]
    accent = RGBColor(0x0A, 0x46, 0x40)

    by_id = {p["id"]: p for p in data["projects"]}
    order = profile.get("include_projects")
    projects = [by_id[i] for i in order if i in by_id] if order else data["projects"]

    doc = Document()
    doc.add_heading(person.get("Name", ""), level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(profile.get("headline") or person.get("Titel/Positionierung", ""))
    run.bold = True
    doc.add_paragraph(person.get("Untertitel", ""))

    # 'Rate Remote'/'Rate Vor-Ort' were replaced by a single 'Tagessatz' on
    # 2026-08-10. The exporter kept asking for the old keys, so the published
    # cv.docx read "Remote 95 % ·  remote /  vor Ort" — two empty slots where
    # the day rate belongs, on the one artifact that goes out by e-mail.
    kond = data["konditionen"]
    facts = [
        f"{L['fact_available']} {kond.get('Verfügbarkeit','')}",
        kond.get("Einsatzort", ""),
        f"{L['fact_remote']} {data['remote_pct']}",
    ]
    rate = profile.get("rate") or kond.get("Tagessatz", "")
    if rate:
        facts.append(f"{profile.get('rate_label') or L['fact_rate']} {rate}")
    doc.add_paragraph(" · ".join(f for f in facts if f))

    def heading(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = accent

    heading(L["sec_highlights"])
    for h in (profile.get("highlights") or data["highlights"]):
        clean = h.replace("**", "")
        doc.add_paragraph(clean, style="List Bullet")

    heading(L["sec_skills"])
    for g in data["skills"]:
        p = doc.add_paragraph()
        p.add_run(f"{g['name']}: ").bold = True
        p.add_run(", ".join(g["tags"]))

    heading(L["sec_projects"])
    for pr in projects:
        head = doc.add_paragraph()
        r = head.add_run(f"{pr['period']} ({pr['dur']}) — {pr['title']}")
        r.bold = True
        meta = " · ".join(x for x in (pr["client"], pr["location"], pr["branch"]) if x)
        doc.add_paragraph(meta)
        if pr["roles"]:
            doc.add_paragraph(f'{L["sec_roles"]}: ' + " · ".join(pr["roles"]))
        doc.add_paragraph(pr["desc"].replace("**", ""))
        if pr["tech"]:
            doc.add_paragraph("Tech: " + ", ".join(pr["tech"]))

    heading(f'{L["sub_education"]} & {L["sub_certificates"]}')
    for e in data["education"] + data["certificates"]:
        p = doc.add_paragraph()
        p.add_run(f"{e['date']} — {e['title']}").bold = True
        p.add_run(f", {e['org']}")

    heading(L["sub_contact"])
    # Field names are the schema and stay German; only the printed label is
    # translated. Website/GitHub/LinkedIn/Kaggle are proper nouns either way.
    for key in ("Telefon", "Website", "GitHub", "LinkedIn", "Kaggle"):
        if person.get(key):
            doc.add_paragraph(f'{L["c_phone"] if key == "Telefon" else key}: {person[key]}')

    path = Path(path)
    doc.save(str(path))
    return path
